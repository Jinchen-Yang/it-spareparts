"""智能体文件原语层。

设计思想（借鉴 Claude Code 等智能体框架）：客户文件格式千变万化——询价单、
整机配置（Word/Excel/PDF/txt/图片）——**不写死解析规则**，给模型"眼睛"
（inspect/read_document：看结构与原样内容，自己判断表头/型号列/拆件）和
"手"（write_excel 回填模板 / write_report 生成美化报表）。

安全边界：不提供任意代码执行（多用户后端 exec = RCE）；上传文件只读，写操作
一律产出新 file_id（绝不改写原上传件）；file_id 白名单正则防路径穿越；
扩展名白名单防可执行文件。
"""
import errno
import hashlib
import json
import math
import os
import re
import stat
import tempfile
import unicodedata
import uuid
import zipfile
from collections.abc import Callable
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta, timezone
from decimal import Decimal
from io import BytesIO
from pathlib import Path, PurePosixPath
from typing import Any, Protocol

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import column_index_from_string, get_column_letter
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.config import get_settings
from app.db import SessionLocal
from app.models.agent_artifact import AgentArtifact, AgentArtifactAudit
from app.models.system import SysUser
from app.services import agent_artifact_provenance as artifact_provenance
from app.services import agent_integrity

_LEGACY_FILE_ID = re.compile(r"^[a-f0-9]{12}$")
_MAX_UPLOAD_MB = 20
_MAX_DOWNLOAD_BYTES = _MAX_UPLOAD_MB * 1024 * 1024
_PREVIEW_ROWS = 8
_PREVIEW_COLS = 12
_MAX_READ_ROWS = 200
_MAX_WRITE_CELLS = 3000
_MAX_REPORT_ROWS = 5000
_MAX_REPORT_COLUMNS = 256
_MAX_REPORT_CELLS = 100_000
_MAX_MONEY_COLUMNS = 256
_MAX_RENDER_TEXT_BYTES = 2 * 1024 * 1024
_MAX_RENDER_VALUE_BYTES = 64 * 1024
_MAX_LEGACY_META_BYTES = 256 * 1024
_MAX_ARTIFACT_JSON_BYTES = 512 * 1024
_MAX_PROVENANCE_DERIVATION_DEPTH = 16
_MAX_PROVENANCE_AUTH_NODES = 256
_MAX_PROVENANCE_AUTH_WORK = 1024
_MAX_PROVENANCE_AUTH_MEMO = 256
_MAX_PROVENANCE_AUTH_PATH = 32
_CELL_TRUNC = 60
_DOC_CHAR_CAP = 60_000          # read_document 文本上限，防超长文件撑爆上下文

# 上传扩展名白名单（小写，无点）
_TEXT_EXT = {"txt", "csv", "md"}
_IMG_EXT = {"jpg", "jpeg", "png", "webp", "bmp"}
_ALLOWED_EXT = {"xlsx", "docx", "pdf"} | _TEXT_EXT | _IMG_EXT
_MIME_BY_EXT = {
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "txt": "text/plain; charset=utf-8",
    "csv": "text/csv; charset=utf-8",
    "md": "text/markdown; charset=utf-8",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "png": "image/png",
    "webp": "image/webp",
    "bmp": "image/bmp",
}
_FORMULA_PREFIXES = ("=", "+", "-", "@")


class FileError(Exception):
    """文件层业务错误（消息可直接回给模型/用户）。"""


class ArtifactV2Disabled(FileError):
    """Stable fail-closed signal used by HTTP and tool adapters."""


class ArtifactUnavailable(FileError):
    """File cannot be served; ``reason_code`` is safe for structured audit."""

    def __init__(self, message: str, reason_code: str):
        super().__init__(message)
        self.reason_code = reason_code


class ProvenanceRequired(FileError):
    """Generated content cannot be published without opaque server evidence."""


class AuthorizationUnavailable(FileError):
    """Live authorization could not be determined; callers must never treat it as allow."""


class ObjectPublicationUncertain(FileError):
    """The final object key may exist after an atomic-link postcondition failed."""


class ArtifactObjectInvalid(FileError):
    """The addressed object is provably absent or violates immutable-store rules."""

    def __init__(self, message: str, reason_code: str):
        super().__init__(message)
        self.reason_code = reason_code


class ArtifactStoreUnavailable(FileError):
    """The object state could not be determined; callers must preserve retry state."""


class ArtifactAuditUnavailable(FileError):
    """A required durable Artifact access fact could not be committed."""


_ARTIFACT_V2_DISABLED_MESSAGE = "Artifact Delivery v2 已停用"


def require_artifact_v2_enabled() -> None:
    if not get_settings().agent_artifact_v2_enabled:
        raise ArtifactV2Disabled(_ARTIFACT_V2_DISABLED_MESSAGE)


def artifact_reason_code(exc: FileError) -> str:
    if isinstance(exc, ArtifactV2Disabled):
        return "v2_disabled"
    if isinstance(exc, ArtifactUnavailable):
        return exc.reason_code
    if isinstance(exc, ProvenanceRequired):
        return "provenance_required"
    if isinstance(exc, (ArtifactStoreUnavailable, AuthorizationUnavailable)):
        return "store_or_authorization_unavailable"
    return "validation_failed"


@dataclass(frozen=True, slots=True)
class VerifiedArtifactOwner:
    """DB-bound principal; every sensitive service call revalidates its live user row."""

    _user_pk: int
    _sub: str
    _token_version: int

    @property
    def sub(self) -> str:
        return self._sub


@dataclass(frozen=True)
class StoredObject:
    path: Path
    size_bytes: int
    sha256: str


@dataclass(frozen=True)
class StoredBytes:
    content: bytes
    size_bytes: int
    sha256: str


@dataclass(frozen=True)
class ArtifactDownload:
    artifact_id: str
    content: bytes
    filename: str
    media_type: str
    size_bytes: int
    sha256: str


class ArtifactStore(Protocol):
    """Artifact object-store seam; the first implementation is local and atomic."""

    def publish_bytes(
        self,
        storage_key: str,
        content: bytes,
        *,
        validator: Callable[[Path], None] | None = None,
    ) -> StoredObject: ...

    def path_for(self, storage_key: str) -> Path: ...

    def inspect(self, storage_key: str) -> StoredObject: ...

    def read_bytes(self, storage_key: str, *, max_bytes: int) -> StoredBytes: ...

    def remove(self, storage_key: str) -> None: ...


class LocalArtifactStore:
    """Filesystem store with same-directory staging and atomic publication."""

    def __init__(self, root: Path):
        requested = Path(os.path.abspath(os.fspath(root)))
        requested.mkdir(parents=True, exist_ok=True)
        try:
            state = requested.lstat()
        except OSError as exc:
            raise ArtifactStoreUnavailable("文件存储根目录暂时不可用") from exc
        if (
            not stat.S_ISDIR(state.st_mode)
            or state.st_uid != os.geteuid()
            or stat.S_IMODE(state.st_mode) & 0o022
        ):
            raise ArtifactStoreUnavailable("文件存储根目录权限或身份无效")
        self.root = requested
        self._root_identity = (state.st_dev, state.st_ino)

    @staticmethod
    def _directory_flags() -> int:
        return (
            os.O_RDONLY
            | getattr(os, "O_DIRECTORY", 0)
            | getattr(os, "O_CLOEXEC", 0)
            | getattr(os, "O_NOFOLLOW", 0)
        )

    def _open_root_fd(self) -> int:
        if not getattr(os, "O_NOFOLLOW", 0):
            raise ArtifactStoreUnavailable("当前平台不支持安全文件存储")
        descriptor: int | None = None
        try:
            descriptor = os.open(self.root, self._directory_flags())
            state = os.fstat(descriptor)
        except OSError as exc:
            if descriptor is not None:
                os.close(descriptor)
            raise ArtifactStoreUnavailable("文件存储根目录暂时不可读取") from exc
        if (
            not stat.S_ISDIR(state.st_mode)
            or (state.st_dev, state.st_ino) != self._root_identity
            or state.st_uid != os.geteuid()
            or stat.S_IMODE(state.st_mode) & 0o022
        ):
            os.close(descriptor)
            raise ArtifactStoreUnavailable("文件存储根目录身份已变化")
        return descriptor

    def _root_path_is_current(self) -> bool:
        try:
            state = self.root.lstat()
        except OSError:
            return False
        return (
            stat.S_ISDIR(state.st_mode)
            and (state.st_dev, state.st_ino) == self._root_identity
            and state.st_uid == os.geteuid()
            and not stat.S_IMODE(state.st_mode) & 0o022
        )

    def _open_owned_directory(self, parent_fd: int, name: str) -> tuple[int, bool]:
        created = False
        descriptor: int | None = None
        try:
            try:
                os.mkdir(name, mode=0o700, dir_fd=parent_fd)
                created = True
            except FileExistsError:
                pass
            descriptor = os.open(name, self._directory_flags(), dir_fd=parent_fd)
            state = os.fstat(descriptor)
        except OSError as exc:
            if descriptor is not None:
                os.close(descriptor)
            raise ArtifactStoreUnavailable("文件存储目录暂时不可用") from exc
        if not stat.S_ISDIR(state.st_mode) or state.st_uid != os.geteuid():
            os.close(descriptor)
            raise ArtifactStoreUnavailable("文件存储目录权限或身份无效")
        try:
            os.fchmod(descriptor, 0o700)
            state = os.fstat(descriptor)
        except OSError as exc:
            os.close(descriptor)
            raise ArtifactStoreUnavailable("文件存储目录权限无法收紧") from exc
        if stat.S_IMODE(state.st_mode) != 0o700:
            os.close(descriptor)
            raise ArtifactStoreUnavailable("文件存储目录权限无效")
        if created:
            try:
                os.fsync(parent_fd)
            except OSError as exc:
                os.close(descriptor)
                raise ArtifactStoreUnavailable("文件存储目录无法持久化") from exc
        return descriptor, created

    @staticmethod
    def _directory_entry_is_current(parent_fd: int, name: str, child_fd: int) -> bool:
        try:
            entry = os.stat(name, dir_fd=parent_fd, follow_symlinks=False)
            opened = os.fstat(child_fd)
        except OSError:
            return False
        return (
            stat.S_ISDIR(entry.st_mode)
            and (entry.st_dev, entry.st_ino) == (opened.st_dev, opened.st_ino)
        )

    def _require_current_directories(
        self, bindings: list[tuple[int, str, int]]
    ) -> None:
        if not self._root_path_is_current() or any(
            not self._directory_entry_is_current(parent, name, child)
            for parent, name, child in bindings
        ):
            raise ArtifactStoreUnavailable("文件存储目录在发布期间发生变化")

    def path_for(self, storage_key: str) -> Path:
        key = str(storage_key or "")
        pure = PurePosixPath(key)
        if not key or pure.is_absolute() or ".." in pure.parts or "\\" in key:
            raise FileError("文件存储定位无效")
        if any(part in {"", ".", ".."} for part in pure.parts):
            raise FileError("文件存储定位无效")
        return self.root.joinpath(*pure.parts)

    def publish_bytes(
        self,
        storage_key: str,
        content: bytes,
        *,
        validator: Callable[[Path], None] | None = None,
    ) -> StoredObject:
        final_path = self.path_for(storage_key)
        pure = PurePosixPath(storage_key)
        directory_fds: list[int] = []
        bindings: list[tuple[int, str, int]] = []
        temp_dir_fd: int | None = None
        final_parent_fd: int | None = None
        temp_fd: int | None = None
        temp_name: str | None = None
        temp_identity: tuple[int, int] | None = None
        linked = False
        try:
            root_fd = self._open_root_fd()
            directory_fds.append(root_fd)
            temp_dir_fd, _ = self._open_owned_directory(root_fd, ".tmp")
            directory_fds.append(temp_dir_fd)
            bindings.append((root_fd, ".tmp", temp_dir_fd))

            current_fd = root_fd
            for part in pure.parts[:-1]:
                child_fd, _ = self._open_owned_directory(current_fd, part)
                directory_fds.append(child_fd)
                bindings.append((current_fd, part, child_fd))
                current_fd = child_fd
            final_parent_fd = current_fd

            temp_name = f"artifact-{uuid.uuid4().hex}.part"
            temp_fd = os.open(
                temp_name,
                os.O_RDWR
                | os.O_CREAT
                | os.O_EXCL
                | getattr(os, "O_CLOEXEC", 0)
                | getattr(os, "O_NOFOLLOW", 0),
                0o600,
                dir_fd=temp_dir_fd,
            )
            os.fchmod(temp_fd, 0o600)
            initial_temp_state = os.fstat(temp_fd)
            temp_identity = (
                initial_temp_state.st_dev,
                initial_temp_state.st_ino,
            )
            digest = hashlib.sha256()
            with os.fdopen(temp_fd, "wb", closefd=False) as handle:
                view = memoryview(content)
                for offset in range(0, len(view), 1024 * 1024):
                    chunk = view[offset:offset + 1024 * 1024]
                    handle.write(chunk)
                    digest.update(chunk)
                handle.flush()
                os.fsync(handle.fileno())
            if validator is not None:
                try:
                    validator(Path("/proc/self/fd") / str(temp_fd))
                except FileError:
                    raise
                except Exception as exc:  # noqa: BLE001 - trusted format validator
                    raise FileError("文件格式校验失败") from exc

            temp_state = os.fstat(temp_fd)
            if (
                not stat.S_ISREG(temp_state.st_mode)
                or stat.S_IMODE(temp_state.st_mode) != 0o600
                or temp_state.st_size != len(content)
            ):
                raise ArtifactStoreUnavailable("文件暂存对象在校验期间发生变化")
            os.lseek(temp_fd, 0, os.SEEK_SET)
            verified_digest = hashlib.sha256()
            verified_size = 0
            while chunk := os.read(temp_fd, 1024 * 1024):
                verified_size += len(chunk)
                verified_digest.update(chunk)
            if (
                verified_size != len(content)
                or verified_digest.digest() != digest.digest()
            ):
                raise ArtifactStoreUnavailable("文件暂存对象完整性校验失败")
            os.fsync(temp_fd)

            temp_entry = os.stat(
                temp_name,
                dir_fd=temp_dir_fd,
                follow_symlinks=False,
            )
            if (
                not stat.S_ISREG(temp_entry.st_mode)
                or (temp_entry.st_dev, temp_entry.st_ino)
                != (temp_state.st_dev, temp_state.st_ino)
            ):
                raise ArtifactStoreUnavailable("文件暂存对象身份已变化")
            self._require_current_directories(bindings)

            # Hard-link publication is atomic and fails if the destination exists.
            # Unlike exists()+replace it cannot overwrite another concurrent winner.
            try:
                os.link(
                    temp_name,
                    pure.parts[-1],
                    src_dir_fd=temp_dir_fd,
                    dst_dir_fd=final_parent_fd,
                    follow_symlinks=False,
                )
            except FileExistsError:
                raise FileError("文件发布冲突，请重试")
            linked = True
            final_fd = os.open(
                pure.parts[-1],
                os.O_RDONLY
                | getattr(os, "O_CLOEXEC", 0)
                | getattr(os, "O_NOFOLLOW", 0),
                dir_fd=final_parent_fd,
            )
            try:
                final_state = os.fstat(final_fd)
            finally:
                os.close(final_fd)
            if (
                not stat.S_ISREG(final_state.st_mode)
                or (final_state.st_dev, final_state.st_ino)
                != (temp_state.st_dev, temp_state.st_ino)
                or final_state.st_size != len(content)
            ):
                raise ObjectPublicationUncertain(
                    "文件对象发布身份无法确认"
                )
            os.fsync(final_parent_fd)
            current_temp_entry = os.stat(
                temp_name,
                dir_fd=temp_dir_fd,
                follow_symlinks=False,
            )
            if (
                temp_identity is None
                or (current_temp_entry.st_dev, current_temp_entry.st_ino)
                != temp_identity
            ):
                raise ObjectPublicationUncertain("文件暂存对象名称已变化")
            os.unlink(temp_name, dir_fd=temp_dir_fd)
            temp_name = None
            if temp_dir_fd != final_parent_fd:
                os.fsync(temp_dir_fd)
            self._require_current_directories(bindings)
            return StoredObject(
                path=final_path,
                size_bytes=final_state.st_size,
                sha256=digest.hexdigest(),
            )
        except OSError as exc:
            if linked:
                raise ObjectPublicationUncertain(
                    "文件对象已原子发布但持久化结果待协调"
                ) from exc
            raise ArtifactStoreUnavailable("文件暂时无法安全发布") from exc
        except Exception as exc:
            if linked:
                # Never unlink final_path here.  After the hard link succeeds, any
                # unlink/fsync/stat failure is an UNKNOWN durable-publication outcome;
                # the validating row is the only safe reconciliation marker.
                raise ObjectPublicationUncertain(
                    "文件对象已原子发布但持久化结果待协调"
                ) from exc
            raise
        finally:
            if temp_name is not None and temp_dir_fd is not None:
                try:
                    entry = os.stat(
                        temp_name,
                        dir_fd=temp_dir_fd,
                        follow_symlinks=False,
                    )
                    if temp_identity is not None and (
                        entry.st_dev, entry.st_ino
                    ) == temp_identity:
                        os.unlink(temp_name, dir_fd=temp_dir_fd)
                except OSError:
                    pass
            if temp_fd is not None:
                os.close(temp_fd)
            for directory_fd in reversed(directory_fds):
                os.close(directory_fd)

    def inspect(self, storage_key: str) -> StoredObject:
        path = self.path_for(storage_key)
        stored = self.read_bytes(storage_key, max_bytes=_MAX_DOWNLOAD_BYTES)
        return StoredObject(
            path=path,
            size_bytes=stored.size_bytes,
            sha256=stored.sha256,
        )

    def read_bytes(self, storage_key: str, *, max_bytes: int) -> StoredBytes:
        """Read/hash from one no-follow regular-file handle under a hard budget."""
        if isinstance(max_bytes, bool) or not isinstance(max_bytes, int) or max_bytes < 0:
            raise FileError("文件读取大小预算无效")
        self.path_for(storage_key)
        pure = PurePosixPath(storage_key)
        directory_flags = os.O_RDONLY | getattr(os, "O_DIRECTORY", 0)
        nofollow = getattr(os, "O_NOFOLLOW", 0)
        cloexec = getattr(os, "O_CLOEXEC", 0)
        directory_fds: list[int] = []
        file_fd: int | None = None
        try:
            try:
                current_fd = self._open_root_fd()
            except OSError as exc:
                raise ArtifactStoreUnavailable(
                    "文件存储根目录暂时不可读取"
                ) from exc
            directory_fds.append(current_fd)
            for part in pure.parts[:-1]:
                try:
                    current_fd = os.open(
                        part,
                        directory_flags | nofollow | cloexec,
                        dir_fd=current_fd,
                    )
                except OSError as exc:
                    # A missing/unmounted parent does not prove that the final
                    # immutable object is absent.  Preserve reconciliation state.
                    raise ArtifactStoreUnavailable(
                        "文件存储目录暂时不可读取"
                    ) from exc
                directory_fds.append(current_fd)
            try:
                file_fd = os.open(
                    pure.parts[-1],
                    os.O_RDONLY | nofollow | cloexec,
                    dir_fd=current_fd,
                )
            except OSError as exc:
                if exc.errno in {errno.ENOENT, errno.ENOTDIR}:
                    raise ArtifactObjectInvalid(
                        "文件对象不存在或定位无效",
                        "object_missing",
                    ) from exc
                if exc.errno == errno.ELOOP:
                    raise ArtifactObjectInvalid(
                        "文件对象定位违反不可变存储规则",
                        "object_invalid",
                    ) from exc
                raise ArtifactStoreUnavailable("文件暂时不可读取") from exc
            before = os.fstat(file_fd)
            if not stat.S_ISREG(before.st_mode):
                raise ArtifactObjectInvalid("文件对象不是常规文件", "object_invalid")
            if before.st_size > max_bytes:
                raise ArtifactObjectInvalid("文件超过下载大小预算", "object_oversize")
            digest = hashlib.sha256()
            chunks: list[bytes] = []
            size = 0
            with os.fdopen(file_fd, "rb", closefd=True) as handle:
                file_fd = None
                for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                    size += len(chunk)
                    if size > max_bytes:
                        raise ArtifactObjectInvalid(
                            "文件超过下载大小预算",
                            "object_oversize",
                        )
                    digest.update(chunk)
                    chunks.append(chunk)
                after = os.fstat(handle.fileno())
            if (
                not stat.S_ISREG(after.st_mode)
                or (before.st_dev, before.st_ino) != (after.st_dev, after.st_ino)
                or before.st_size != after.st_size
                or before.st_mtime_ns != after.st_mtime_ns
                or before.st_ctime_ns != after.st_ctime_ns
                or size != after.st_size
            ):
                raise ArtifactStoreUnavailable("文件读取期间发生变化")
        except OSError as exc:
            raise ArtifactStoreUnavailable("文件暂时不可读取") from exc
        finally:
            if file_fd is not None:
                os.close(file_fd)
            for directory_fd in reversed(directory_fds):
                os.close(directory_fd)
        return StoredBytes(
            content=b"".join(chunks),
            size_bytes=size,
            sha256=digest.hexdigest(),
        )

    def remove(self, storage_key: str) -> None:
        self.path_for(storage_key)
        raise ArtifactStoreUnavailable(
            "本地 Artifact Store 未启用条件物理删除"
        )


def _dir() -> Path:
    # 必须落在持久卷内：生产的持久卷挂在 raw_file_dir(/app/data/raw)，放它的【子目录】才不会
    # 随容器重建被清空——否则每次部署/重启都会清掉 AI 生成的报价单，下载链接全部 404。
    # raw 导入文件是 raw_file_dir/{hash}.xlsx 平铺，agent_files 子目录与之不冲突。
    d = Path(get_settings().raw_file_dir) / "agent_files"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _meta_path(file_id: str) -> Path:
    return _dir() / f"{file_id}.meta.json"


def _ext_of(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _safe_filename(filename: str, default: str = "artifact") -> str:
    """Keep readable Unicode while removing path/header/control ambiguity."""
    raw = unicodedata.normalize("NFKC", str(filename or default))
    raw = "".join(" " if unicodedata.category(ch).startswith("C") else ch for ch in raw)
    raw = re.sub(r"[\\/:*?\"<>|]", "_", raw)
    raw = re.sub(r"\s+", " ", raw).strip().lstrip(".").strip()
    if not raw:
        raw = default
    if len(raw) > 180:
        suffix = Path(raw).suffix[:20]
        raw = f"{raw[:180 - len(suffix)].rstrip()}{suffix}"
    return raw


def _is_legacy_id(file_id: str) -> bool:
    return bool(_LEGACY_FILE_ID.fullmatch(file_id))


def _storage_key(file_id: str, ext: str) -> str:
    if ext not in _ALLOWED_EXT:
        raise FileError("文件类型元数据无效")
    if _is_legacy_id(file_id):
        raise FileError("新制品标识无效")
    return f"objects/{file_id}.{ext}"


def get_artifact_store() -> ArtifactStore:
    """Resolve the active store; kept as a seam for object-store replacement/tests."""
    return LocalArtifactStore(_dir())


def _data_path(file_id: str, ext: str) -> Path:
    fid = _check_id(file_id)
    if ext not in _ALLOWED_EXT:
        raise FileError("文件类型元数据无效")
    if _is_legacy_id(fid):
        return _dir() / f"{fid}.{ext}"
    return get_artifact_store().path_for(_storage_key(fid, ext))


def _check_id(file_id: str) -> str:
    fid = str(file_id or "").strip().lower()
    if _is_legacy_id(fid):
        return fid
    try:
        parsed = uuid.UUID(fid)
    except (ValueError, AttributeError) as exc:
        raise FileError("非法 file_id") from exc
    canonical = str(parsed)
    if canonical != fid:
        raise FileError("非法 file_id")
    return canonical


def _query_artifact(db, file_id: str) -> AgentArtifact | None:
    if _is_legacy_id(file_id):
        return None
    return db.get(AgentArtifact, file_id)


def _bounded_json_object(value: Any, label: str) -> dict[str, Any]:
    if not isinstance(value, dict):
        raise FileError(f"{label} 必须是 JSON 对象")
    try:
        encoded = agent_integrity.canonicalize(value)
        if len(encoded) > _MAX_ARTIFACT_JSON_BYTES:
            raise FileError(f"{label} 超过大小预算")
        copied = json.loads(encoded.decode("utf-8"))
    except FileError:
        raise
    except (agent_integrity.IntegrityError, UnicodeError, json.JSONDecodeError) as exc:
        raise FileError(f"{label} 不是有效的有界 JSON 对象") from exc
    if not isinstance(copied, dict):
        raise FileError(f"{label} 必须是 JSON 对象")
    return copied


def _add_artifact_audit(
    db: Session,
    *,
    artifact_id: str | None,
    action: str,
    outcome: str,
    actor: str,
    decision_key: str | None = None,
    from_status: str | None = None,
    to_status: str | None = None,
    detail: dict[str, Any] | None = None,
) -> None:
    db.add(AgentArtifactAudit(
        artifact_id=artifact_id,
        decision_key=decision_key,
        action=action,
        outcome=outcome,
        actor=actor,
        from_status=from_status,
        to_status=to_status,
        detail=_bounded_json_object(detail or {}, "Artifact 审计 detail"),
    ))


def record_artifact_http_access(
    *,
    action: str,
    outcome: str,
    actor: str,
    artifact_id: str | None = None,
    attempted_identifier: str | None = None,
    reason_code: str | None = None,
    size_bytes: int | None = None,
) -> bool:
    """Commit a specialized HTTP access fact, or decline unknown identifiers.

    Success always requires an existing UUID-backed Artifact.  Denials are durable
    only when they bind to an existing UUID row or the explicitly recognizable
    legacy format.  Malformed/unknown identifiers stay in the generic security log.
    """
    if action not in {"upload", "download", "preview"}:
        raise ArtifactAuditUnavailable("Artifact 访问审计动作无效")
    if outcome not in {"success", "denied"}:
        raise ArtifactAuditUnavailable("Artifact 访问审计结果无效")
    checked_actor = str(actor or "").strip()
    if not checked_actor or len(checked_actor) > 64:
        checked_actor = "unknown:http-principal"
    checked_reason = str(reason_code or "").strip().lower()
    if not re.fullmatch(r"[a-z0-9_]{1,64}", checked_reason):
        checked_reason = "validation_failed"
    if (
        size_bytes is not None
        and (isinstance(size_bytes, bool) or not isinstance(size_bytes, int) or size_bytes < 0)
    ):
        raise ArtifactAuditUnavailable("Artifact 访问审计大小无效")

    target_id: str | None = None
    legacy_denial = False
    if outcome == "success":
        try:
            target_id = _check_id(artifact_id or "")
        except FileError as exc:
            raise ArtifactAuditUnavailable("Artifact 成功审计缺少有效绑定") from exc
        if _is_legacy_id(target_id):
            raise ArtifactAuditUnavailable("Legacy 文件不能产生成功审计")
    else:
        candidate = str(attempted_identifier or "").strip().lower()
        if _is_legacy_id(candidate):
            legacy_denial = True
        else:
            try:
                target_id = _check_id(candidate)
            except FileError:
                return False

    detail: dict[str, Any] = {}
    if outcome == "denied":
        detail["reason_code"] = checked_reason
        detail["identifier_format"] = "legacy" if legacy_denial else "uuid"
    if size_bytes is not None:
        detail["size_bytes"] = size_bytes
    try:
        with SessionLocal.begin() as db:
            if target_id is not None and db.get(AgentArtifact, target_id) is None:
                if outcome == "success":
                    raise ArtifactAuditUnavailable(
                        "Artifact 成功审计无法绑定当前制品"
                    )
                return False
            _add_artifact_audit(
                db,
                artifact_id=target_id,
                action=f"http_{action}",
                outcome=outcome,
                actor=checked_actor,
                detail=detail,
            )
    except ArtifactAuditUnavailable:
        raise
    except Exception as exc:  # noqa: BLE001 - access delivery must fail closed
        raise ArtifactAuditUnavailable("Artifact 访问审计暂时不可用") from exc
    return True


def _require_artifact_binding(meta: dict[str, Any]) -> None:
    try:
        artifact_provenance.verify_artifact_binding(
            meta.get("binding_envelope"),
            meta,
        )
    except artifact_provenance.ProvenanceError as exc:
        raise ArtifactUnavailable(
            "Artifact 元数据绑定校验失败",
            "binding_invalid",
        ) from exc


def _binding_metadata_from_row(row: AgentArtifact) -> dict[str, Any]:
    return {
        "file_id": row.id,
        "owner_sub": row.owner_sub,
        "filename": row.filename,
        "media_type": row.media_type,
        "size_bytes": row.size_bytes,
        "sha256": row.sha256,
        "storage_key": row.storage_key,
        "kind": row.kind,
        "status": row.status,
        "sensitivity": row.sensitivity,
        "source_ids": row.source_ids,
        "access_scope": row.access_scope,
        "extra_meta": row.extra_meta,
        "binding_envelope": row.binding_envelope,
        "created_at": row.created_at,
        "expires_at": row.expires_at,
    }


def _locked_artifact(db: Session, artifact_id: str) -> AgentArtifact | None:
    return db.scalar(
        select(AgentArtifact)
        .where(AgentArtifact.id == artifact_id)
        .with_for_update()
        .execution_options(populate_existing=True)
    )


def _transition_locked_bound_status(
    db: Session,
    row: AgentArtifact,
    *,
    expected: str,
    target: str,
    actor: str,
    reason: str,
) -> bool:
    """Transition one locked row with old proof, new proof, and durable audit."""
    allowed_edges = {
        ("prepared", "validating"),
        ("prepared", "failed"),
        ("validating", "ready"),
        ("validating", "failed"),
        ("ready", "expired"),
    }
    if (expected, target) not in allowed_edges:
        raise FileError("Artifact 状态迁移边无效")
    if row.status != expected:
        return False
    _require_artifact_binding(_binding_metadata_from_row(row))
    row.status = target
    try:
        row.binding_envelope = artifact_provenance.seal_artifact_binding(
            _binding_metadata_from_row(row)
        )
    except artifact_provenance.ProvenanceError as exc:
        row.status = expected
        raise ArtifactUnavailable(
            "Artifact 状态绑定签发失败",
            "binding_resign_failed",
        ) from exc
    _add_artifact_audit(
        db,
        artifact_id=row.id,
        action="status_transition",
        outcome="success",
        actor=actor,
        from_status=expected,
        to_status=target,
        detail={"reason": reason},
    )
    return True


def _validated_artifact_metadata(row: AgentArtifact) -> dict[str, Any]:
    try:
        extra = _bounded_json_object(row.extra_meta, "Artifact extra_meta")
        access_scope = _bounded_json_object(
            row.access_scope,
            "Artifact access_scope",
        )
        binding_envelope = _bounded_json_object(
            row.binding_envelope,
            "Artifact binding_envelope",
        )
        source_ids = row.source_ids
        if (
            not isinstance(source_ids, list)
            or any(
                not isinstance(source_id, str)
                or str(uuid.UUID(source_id)) != source_id
                for source_id in source_ids
            )
            or len(source_ids) != len(set(source_ids))
        ):
            raise FileError("Artifact source_ids 类型无效")
    except (FileError, ValueError, AttributeError) as exc:
        raise ArtifactUnavailable(
            "文件元数据校验失败",
            "metadata_invalid",
        ) from exc
    meta = {
        **extra,
        "file_id": row.id,
        "filename": row.filename,
        "ext": _ext_of(row.filename),
        "kind": row.kind,
        "sensitivity": row.sensitivity,
        "operated_by": row.owner_sub,
        "owner_sub": row.owner_sub,
        "media_type": row.media_type,
        "size_bytes": row.size_bytes,
        "sha256": row.sha256,
        "status": row.status,
        "storage_key": row.storage_key,
        "source_ids": list(source_ids),
        "access_scope": access_scope,
        "extra_meta": extra,
        "binding_envelope": binding_envelope,
        "created_at": row.created_at,
        "expires_at": row.expires_at,
    }
    _require_artifact_binding(meta)
    return meta


def _find_artifact_meta(file_id: str, *, require_ready: bool) -> dict | None:
    if _is_legacy_id(file_id):
        return None
    require_artifact_v2_enabled()
    now = datetime.now(timezone.utc)
    with SessionLocal() as db:
        row = _query_artifact(db, file_id)
        if row is None:
            return None
        meta = _validated_artifact_metadata(row)
        if row.expires_at <= now and row.status == "ready":
            # The initial read is only a candidate.  Reload and lock the current row
            # before mutating so stale ORM state can never authorize a transition.
            row = _locked_artifact(db, row.id)
            if row is None:
                return None
            meta = _validated_artifact_metadata(row)
            if row.expires_at <= now and row.status == "ready":
                _transition_locked_bound_status(
                    db,
                    row,
                    expected="ready",
                    target="expired",
                    actor="system:artifact-read-expiry",
                    reason="retention_expired",
                )
            try:
                db.commit()
            except Exception as exc:  # noqa: BLE001 - fail closed if audit is not durable
                db.rollback()
                raise ArtifactUnavailable(
                    "文件状态暂时无法确认", "state_audit_failed"
                ) from exc
            meta = _validated_artifact_metadata(row)
        if require_ready and meta["status"] != "ready":
            reason = (
                meta["status"]
                if meta["status"] in {"expired", "failed"}
                else "not_ready"
            )
            raise ArtifactUnavailable("文件不存在或不可下载", reason)
        return meta


def _artifact_meta(file_id: str, *, require_ready: bool) -> dict:
    meta = _find_artifact_meta(file_id, require_ready=require_ready)
    if meta is None:
        raise ArtifactUnavailable("文件不存在或已清理", "not_found")
    return meta


def _read_verified_artifact(
    meta: dict[str, Any],
    *,
    require_media_type: bool = True,
) -> StoredBytes:
    """Return the exact bytes verified against immutable Artifact metadata."""
    fid = meta["file_id"]
    ext = meta.get("ext", "")
    if meta.get("filename") != _safe_filename(meta.get("filename", "")):
        raise FileError("文件元数据校验失败")
    if require_media_type and meta.get("media_type") != _MIME_BY_EXT.get(ext):
        raise FileError("文件元数据校验失败")
    expected_key = _storage_key(fid, ext)
    if meta.get("storage_key") != expected_key:
        raise FileError("文件元数据校验失败")
    try:
        stored = get_artifact_store().read_bytes(
            expected_key,
            max_bytes=_MAX_DOWNLOAD_BYTES,
        )
    except ArtifactStoreUnavailable:
        raise
    except ArtifactObjectInvalid as exc:
        reason = (
            "size_limit"
            if exc.reason_code == "object_oversize"
            else exc.reason_code
        )
        message = (
            "文件超过允许的下载大小"
            if reason == "size_limit"
            else "文件对象不存在、不可用或已清理"
        )
        raise ArtifactUnavailable(message, reason) from exc
    except FileError as exc:
        raise ArtifactStoreUnavailable("文件对象状态暂时无法确认") from exc
    if stored.size_bytes != meta.get("size_bytes") or stored.sha256 != meta.get("sha256"):
        raise ArtifactUnavailable("文件完整性校验失败", "integrity_failed")
    return stored


def _verify_artifact(meta: dict[str, Any]) -> None:
    """Compatibility verifier for authorization/reconciliation paths."""
    _read_verified_artifact(meta)


def _load_verified_artifact(file_id: str) -> tuple[dict[str, Any], StoredBytes]:
    fid = _check_id(file_id)
    if _is_legacy_id(fid):
        raise ArtifactUnavailable("文件不存在或无权访问", "not_found_or_forbidden")
    meta = _artifact_meta(fid, require_ready=True)
    return meta, _read_verified_artifact(meta)


def _safe_download_media_type(meta: dict[str, Any]) -> str:
    """Bind an allowlisted MIME to the server-owned storage extension."""
    ext = _ext_of(meta.get("filename", ""))
    expected = _MIME_BY_EXT.get(ext)
    if expected is None or meta.get("media_type") != expected:
        return "application/octet-stream"
    return expected


def _validate_legacy_meta(value: Any) -> dict[str, Any]:
    """Validate the complete historical sidecar shape before any projection.

    Sidecars predate authenticated Artifact ownership and remain untrusted.  Parsing
    them is retained only for conservative offline inspection; no field in this
    object can mint a current owner capability.
    """
    required = {"filename", "ext", "kind", "operated_by", "created_at"}
    optional = {"sheets", "base_file_id", "report"}
    if not isinstance(value, dict) or not required <= set(value) <= required | optional:
        raise FileError("文件元数据损坏")
    filename = value["filename"]
    ext = value["ext"]
    kind = value["kind"]
    operated_by = value["operated_by"]
    created_at = value["created_at"]
    if (
        not isinstance(filename, str)
        or not 0 < len(filename.encode("utf-8")) <= 4096
        or not isinstance(ext, str)
        or ext not in _ALLOWED_EXT
        or not isinstance(kind, str)
        or kind not in {"upload", "generated"}
        or (operated_by is not None and (
            not isinstance(operated_by, str)
            or not 0 < len(operated_by.encode("utf-8")) <= 256
        ))
        or not isinstance(created_at, str)
        or not 0 < len(created_at) <= 64
    ):
        raise FileError("文件元数据损坏")
    try:
        parsed_created_at = datetime.fromisoformat(created_at)
    except ValueError as exc:
        raise FileError("文件元数据损坏") from exc
    if parsed_created_at.tzinfo is None:
        raise FileError("文件元数据损坏")

    if "sheets" in value:
        sheets = value["sheets"]
        if not isinstance(sheets, list) or len(sheets) > 64:
            raise FileError("文件元数据损坏")
        for sheet in sheets:
            if not isinstance(sheet, dict) or set(sheet) != {"name", "n_rows", "n_cols"}:
                raise FileError("文件元数据损坏")
            name = sheet["name"]
            n_rows = sheet["n_rows"]
            n_cols = sheet["n_cols"]
            if (
                not isinstance(name, str)
                or not 0 < len(name) <= 31
                or isinstance(n_rows, bool)
                or not isinstance(n_rows, int)
                or not 0 <= n_rows <= 1_048_576
                or isinstance(n_cols, bool)
                or not isinstance(n_cols, int)
                or not 0 <= n_cols <= 16_384
            ):
                raise FileError("文件元数据损坏")
    if "base_file_id" in value:
        base_file_id = value["base_file_id"]
        if base_file_id is not None:
            if not isinstance(base_file_id, str):
                raise FileError("文件元数据损坏")
            try:
                _check_id(base_file_id)
            except FileError as exc:
                raise FileError("文件元数据损坏") from exc
    if "report" in value and not isinstance(value["report"], bool):
        raise FileError("文件元数据损坏")
    return dict(value)


def _load_meta(file_id: str) -> dict:
    fid = _check_id(file_id)
    meta = _find_artifact_meta(fid, require_ready=True)
    if meta is not None:
        _verify_artifact(meta)
        return meta
    if not _is_legacy_id(fid):
        raise ArtifactUnavailable("文件不存在或已清理", "not_found")
    p = _meta_path(fid)
    if not p.exists():
        raise ArtifactUnavailable("文件不存在或已清理", "not_found")
    try:
        if p.stat().st_size > _MAX_LEGACY_META_BYTES:
            raise FileError("文件元数据损坏")
        meta = _validate_legacy_meta(json.loads(p.read_text(encoding="utf-8")))
    except FileError:
        raise
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        raise FileError("文件元数据损坏") from exc
    ext = meta["ext"]
    if ext not in _ALLOWED_EXT or not _data_path(fid, ext).is_file():
        raise ArtifactUnavailable("文件不存在或已清理", "object_missing")
    meta["filename"] = _safe_filename(meta.get("filename", f"{fid}.{ext}"))
    meta["media_type"] = _MIME_BY_EXT[ext]
    return meta


def _save_meta(file_id: str, meta: dict) -> None:
    """Legacy writes are permanently disabled; retained only as a deny canary."""
    raise ArtifactUnavailable("Legacy 文件元数据写入已停用", "legacy_denied")


def artifact_info(file_id: str) -> dict:
    """Return structured, non-path metadata for a new Artifact."""
    fid = _check_id(file_id)
    if _is_legacy_id(fid):
        raise ArtifactUnavailable("Legacy 文件不可访问", "legacy_denied")
    db_meta = _find_artifact_meta(fid, require_ready=False)
    if db_meta is not None:
        return {
            key: value.isoformat() if isinstance(value, datetime) else value
            for key, value in db_meta.items()
            if key in {
                "file_id", "filename", "media_type", "size_bytes", "sha256",
                "status", "source_ids", "created_at", "expires_at", "kind",
                "sensitivity",
            }
        }
    raise FileError("文件不存在或已清理")


def _artifact_ref(info: dict) -> dict:
    """Stable transport shape for tools/SSE/message persistence consumers."""
    artifact_id = info["file_id"]
    return {
        "id": artifact_id,
        "filename": info["filename"],
        "mime_type": info["media_type"],
        "size_bytes": info["size_bytes"],
        "sha256": info["sha256"],
        "status": info["status"],
        "sensitivity": info["sensitivity"],
        "download_url": f"/api/agent/files/{artifact_id}",
    }


def _validate_staged_file(path: Path, ext: str) -> None:
    if ext == "xlsx":
        with path.open("rb") as handle:
            workbook = load_workbook(handle, read_only=True, data_only=True)
            workbook.close()
        return
    if ext == "docx":
        if not zipfile.is_zipfile(path):
            raise FileError("无法解析 docx（文件容器损坏）")
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise FileError("无法解析 docx（缺少 Word 文档结构）")
        return
    if ext == "pdf":
        with path.open("rb") as handle:
            if handle.read(5) != b"%PDF-":
                raise FileError("无法解析 pdf（文件签名不匹配）")
        return
    if ext in _IMG_EXT:
        from PIL import Image

        expected = "JPEG" if ext in {"jpg", "jpeg"} else ext.upper()
        with Image.open(path) as image:
            image.verify()
            if image.format != expected:
                raise FileError("图片扩展名与实际格式不一致")
        return
    if ext in _TEXT_EXT:
        data = path.read_bytes()
        if b"\x00" in data:
            raise FileError("文本文件包含二进制内容")
        try:
            data.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise FileError("文本文件必须使用 UTF-8 编码") from exc


def snapshot_access_scope(user_ctx: Any) -> dict:
    """Capture server-authenticated visibility; never accept this from model arguments."""
    from app import permissions

    stable_owner_sub(user_ctx)
    permission_source = (
        permissions.effective(user_ctx.role, None)
        if user_ctx.permissions is None
        else user_ctx.permissions
    )
    effective = permissions.runtime_safe(permission_source)
    data_permissions = {key: bool(effective.get(key, False)) for key in permissions.DATA_GROUPS}
    page_permissions = {key: bool(effective.get(key, False)) for key in permissions.PAGE_KEYS}
    visible_field_groups = sorted({
        group
        for key, groups in permissions.DATA_GROUPS.items()
        if data_permissions[key]
        for group in groups
    })
    required = sorted(
        key for key, enabled in effective.items()
        if enabled and (key.startswith("data_") or key.startswith("page_"))
    )
    own_customers = bool(effective.get("own_customers_only"))
    row_subject = (
        _canonical_salesperson_subject(user_ctx.salesperson_name)
        if own_customers
        else None
    )
    if own_customers and not row_subject:
        raise FileError("本人客户范围制品缺少已绑定的销售主体")
    return {
        "version": 1,
        "policy": "current_scope_dominates",
        "required_permissions": required,
        "data_permissions": data_permissions,
        "page_permissions": page_permissions,
        "visible_field_groups": visible_field_groups,
        "row_scope": "own_customers" if own_customers else "all",
        "row_subject": row_subject,
    }


def _default_access_scope(kind: str) -> dict:
    if kind == "upload":
        return artifact_provenance.classify_upload_access_scope(b"", "")
    # Generated output without a server-authenticated snapshot is explicitly unclassified.
    return {
        "version": 1,
        "policy": "unclassified_deny",
        "required_permissions": [],
    }


def _derive_sensitivity(kind: str, access_scope: dict) -> str:
    """Derive classification server-side; callers cannot lower it through tool arguments."""
    if kind == "upload":
        return (
            access_scope.get("sensitivity", "critical")
            if access_scope.get("schema_version") == artifact_provenance.ACCESS_SCHEMA_VERSION
            else "critical"
        )
    if access_scope.get("schema_version") == artifact_provenance.ACCESS_SCHEMA_VERSION:
        return artifact_provenance.derive_sensitivity(
            access_scope.get("required_permissions", [])
        )
    if access_scope.get("policy") != "current_scope_dominates":
        return "critical"
    from app import permissions

    rank = {"low": 0, "medium": 1, "high": 2, "critical": 3}
    required = access_scope.get("required_permissions")
    if not isinstance(required, list):
        return "critical"
    levels = [
        permissions.PERMISSION_META.get(key, {}).get("sensitivity")
        for key in required
    ]
    known = [level for level in levels if level in rank]
    return max(known, key=rank.__getitem__) if known else "high"


def stable_owner_sub(user_ctx: Any) -> str:
    """Return a non-spoofable owner subject or fail closed for shared/guest identities."""
    subject = str(getattr(user_ctx, "user_id", None) or "").strip()
    if (
        not getattr(user_ctx, "is_authenticated", False)
        or getattr(user_ctx, "authn", None) != "sys_user"
        or not getattr(user_ctx, "has_stable_subject", False)
        or not subject
    ):
        raise FileError("创建或访问制品需要实名系统账号")
    return subject


def _context_from_user(user: SysUser):
    from app import permissions
    from app.security import UserContext

    return UserContext(
        user_id=user.username,
        role=user.role,
        salesperson_name=user.salesperson_name,
        permissions=permissions.runtime_safe(permissions.effective_for_user(user)),
        is_authenticated=True,
        authn="sys_user",
        has_stable_subject=True,
        token_version=int(user.token_version or 0),
    )


def verified_artifact_owner(db: Session, user_ctx: Any) -> VerifiedArtifactOwner:
    """Mint a principal only when token claims match the active SysUser row now."""
    from app import permissions

    subject = stable_owner_sub(user_ctx)
    user = db.scalar(
        select(SysUser)
        .where(SysUser.username == subject, SysUser.is_active.is_(True))
        .execution_options(populate_existing=True)
    )
    if user is None:
        raise FileError("创建或访问制品需要有效的当前登录状态")
    current = _context_from_user(user)
    claimed_permissions = getattr(user_ctx, "permissions", None)
    claims_match = (
        getattr(user_ctx, "token_version", None) == current.token_version
        and getattr(user_ctx, "role", None) == current.role
        and _canonical_salesperson_subject(getattr(user_ctx, "salesperson_name", None))
        == _canonical_salesperson_subject(current.salesperson_name)
        and isinstance(claimed_permissions, dict)
        and permissions.runtime_safe(claimed_permissions) == current.permissions
    )
    if not claims_match:
        raise FileError("登录声明与当前账号事实不一致")
    return VerifiedArtifactOwner(user.id, subject, current.token_version or 0)


def _verified_owner_context(owner: VerifiedArtifactOwner):
    if not isinstance(owner, VerifiedArtifactOwner):
        raise FileError("创建或访问制品需要已验证身份")
    with SessionLocal() as db:
        user = db.scalar(select(SysUser).where(
            SysUser.id == owner._user_pk,
            SysUser.username == owner.sub,
            SysUser.is_active.is_(True),
        ))
        if user is None or int(user.token_version or 0) != owner._token_version:
            raise FileError("创建或访问制品需要有效的当前登录状态")
        return _context_from_user(user)


def _verified_owner_sub(owner: VerifiedArtifactOwner) -> str:
    return stable_owner_sub(_verified_owner_context(owner))


def _require_live_owner(owner: VerifiedArtifactOwner) -> str:
    """Classify unexpected principal lookup failures as retryable uncertainty."""
    try:
        return _verified_owner_sub(owner)
    except FileError:
        raise
    except Exception as exc:  # noqa: BLE001 - publisher must preserve retry state
        raise AuthorizationUnavailable("当前账号状态暂时无法确认") from exc


def _generated_scope(owner: VerifiedArtifactOwner) -> dict:
    return snapshot_access_scope(_verified_owner_context(owner))


def _canonical_source_id(source_id: str, owner: VerifiedArtifactOwner) -> str:
    try:
        checked = _authorized_owner_id(source_id, owner)
    except ArtifactUnavailable as exc:
        raise ArtifactUnavailable("无权引用来源制品", exc.reason_code) from exc
    meta = _load_meta(checked)
    return meta.get("file_id", checked)


def _mark_artifact_ready(artifact_id: str, stored: StoredObject) -> None:
    """Lock, re-prove, and durably transition ``validating`` to ``ready``."""
    with SessionLocal.begin() as db:
        row = _locked_artifact(db, artifact_id)
        if row is None:
            raise FileError("文件发布状态冲突")
        meta = _validated_artifact_metadata(row)
        metadata_matches = (
            row.size_bytes == stored.size_bytes and row.sha256 == stored.sha256
        )
        if not metadata_matches:
            raise FileError("文件发布完整性校验失败")
        if row.status == "ready":
            # Idempotency is safe only when the current sealed row and final object
            # both still match and live authorization still holds.  A concurrent
            # winner does not let this publication invocation bypass a later revoke.
            _read_verified_artifact(meta)
            if not _reconcile_ready_authorized(row):
                raise FileError("文件发布授权已失效")
            return
        if row.status != "validating":
            raise FileError("文件发布状态冲突")
        current = _read_verified_artifact(meta)
        if current.size_bytes != stored.size_bytes or current.sha256 != stored.sha256:
            raise FileError("文件发布完整性校验失败")
        if not _reconcile_ready_authorized(row):
            raise FileError("文件发布授权已失效")
        _transition_locked_bound_status(
            db,
            row,
            expected="validating",
            target="ready",
            actor=row.owner_sub,
            reason="publisher_verified",
        )


def _mark_artifact_validating(artifact_id: str) -> None:
    with SessionLocal.begin() as db:
        row = _locked_artifact(db, artifact_id)
        if row is None or row.status != "prepared":
            raise FileError("文件发布状态冲突")
        _transition_locked_bound_status(
            db,
            row,
            expected="prepared",
            target="validating",
            actor=row.owner_sub,
            reason="publisher_started",
        )


def _mark_artifact_failed(artifact_id: str) -> None:
    with SessionLocal.begin() as db:
        row = _locked_artifact(db, artifact_id)
        if row is not None and row.status in {"prepared", "validating"}:
            _transition_locked_bound_status(
                db,
                row,
                expected=row.status,
                target="failed",
                actor=row.owner_sub,
                reason="publisher_failed",
            )


def _publish_artifact(
    content: bytes,
    filename: str,
    *,
    kind: str,
    owner: VerifiedArtifactOwner,
    source_ids: list[str] | None = None,
    extra_meta: dict | None = None,
    provenance_scope: dict | None = None,
) -> dict:
    require_artifact_v2_enabled()
    safe_name = _safe_filename(filename)
    ext = _ext_of(safe_name)
    if ext not in _ALLOWED_EXT:
        raise FileError("文件类型不受支持")
    artifact_id = str(uuid.uuid4())
    storage_key = _storage_key(artifact_id, ext)
    created_at = datetime.now(timezone.utc)
    expires_at = created_at + timedelta(days=get_settings().agent_artifact_retention_days)
    expected_hash = hashlib.sha256(content).hexdigest()
    owner_sub = _require_live_owner(owner)
    provided_sources = (
        [_canonical_source_id(source_id, owner) for source_id in source_ids]
        if source_ids is not None
        else None
    )
    if kind == "upload":
        resolved_scope = _bounded_json_object(
            provenance_scope or _default_access_scope(kind),
            "Artifact access_scope",
        )
    elif kind == "generated":
        if provenance_scope is None:
            raise ProvenanceRequired("生成文件缺少可验证的数据来源，已停止发布")
        resolved_scope = _bounded_json_object(
            provenance_scope, "Artifact access_scope"
        )
        try:
            proven_sources = artifact_provenance.source_artifact_ids(resolved_scope)
        except artifact_provenance.ProvenanceError as exc:
            raise ProvenanceRequired("生成文件来源快照无效，已停止发布") from exc
        if provided_sources is not None and provided_sources != proven_sources:
            raise ProvenanceRequired("生成文件来源列表与证明不一致，已停止发布")
        # Renderer work may be long-running.  Treat the signed snapshot as evidence,
        # not a lease: reload the current principal and every live source before the
        # prepared row is created.  The state machine repeats this at both object-store
        # and ready-transition boundaries below.
        sources = _reauthorize_provenance_scope(owner, resolved_scope)
        if sources != proven_sources:
            raise ProvenanceRequired("生成文件来源列表与当前授权不一致，已停止发布")
    else:
        raise FileError("文件类型不受支持")
    if kind == "upload":
        if provided_sources:
            raise FileError("上传件不能伪造派生来源")
        sources = []
    if extra_meta is not None and not isinstance(extra_meta, dict):
        raise FileError("Artifact extra_meta 必须是 JSON 对象")
    persisted_extra_meta = _bounded_json_object(
        extra_meta or {}, "Artifact extra_meta"
    )
    # A crash reconciler must reconstruct the exact principal capability that began
    # publication; owner_sub alone cannot detect logout/revocation via token_version.
    persisted_extra_meta["_publisher_user_id"] = owner._user_pk
    persisted_extra_meta["_publisher_token_version"] = owner._token_version
    sensitivity = _derive_sensitivity(kind, resolved_scope)
    binding_metadata = {
        "file_id": artifact_id,
        "owner_sub": owner_sub,
        "filename": safe_name,
        "media_type": _MIME_BY_EXT[ext],
        "size_bytes": len(content),
        "sha256": expected_hash,
        "storage_key": storage_key,
        "kind": kind,
        "status": "prepared",
        "sensitivity": sensitivity,
        "source_ids": list(sources),
        "access_scope": resolved_scope,
        "extra_meta": persisted_extra_meta,
        "created_at": created_at,
        "expires_at": expires_at,
    }
    try:
        binding_envelope = artifact_provenance.seal_artifact_binding(
            binding_metadata
        )
    except artifact_provenance.ProvenanceError as exc:
        raise ProvenanceRequired(
            "Artifact 聚合绑定签发失败，已停止发布"
        ) from exc
    with SessionLocal.begin() as db:
        db.add(AgentArtifact(
            id=artifact_id,
            owner_sub=owner_sub,
            filename=safe_name,
            media_type=_MIME_BY_EXT[ext],
            size_bytes=len(content),
            sha256=expected_hash,
            status="prepared",
            storage_key=storage_key,
            kind=kind,
            sensitivity=sensitivity,
            source_ids=sources,
            access_scope=resolved_scope,
            extra_meta=persisted_extra_meta,
            binding_envelope=binding_envelope,
            created_at=created_at,
            expires_at=expires_at,
        ))
        _add_artifact_audit(
            db,
            artifact_id=artifact_id,
            action="artifact_created",
            outcome="success",
            actor=owner_sub,
            to_status="prepared",
            detail={"kind": kind},
        )

    store = get_artifact_store()
    object_published = False
    try:
        _mark_artifact_validating(artifact_id)
        # Uploads also have an external storage side effect.  Revalidate the exact
        # principal capability for every kind after the visible state transition.
        _require_live_owner(owner)
        if kind == "generated":
            # The prepared/validating transition is observable and may race an account
            # change.  Recheck after that transition and immediately before delegating
            # any source-derived bytes to the object store.
            _reauthorize_provenance_scope(owner, resolved_scope)
        stored = store.publish_bytes(
            storage_key,
            content,
            validator=lambda path: _validate_staged_file(path, ext),
        )
        object_published = True
        if stored.size_bytes != len(content) or stored.sha256 != expected_hash:
            raise FileError("文件发布完整性校验失败")
        _require_live_owner(owner)
        if kind == "generated":
            # Object publication is atomic but cannot share a transaction with account
            # and source rows.  A revocation during the write therefore leaves a
            # non-ready reconciliation marker; it never authorizes the ready CAS.
            _reauthorize_provenance_scope(owner, resolved_scope)
        # Re-read the formal key.  A reconciler may win this CAS only after applying the
        # same live principal/generated-scope contract; an identical ready row is
        # idempotent.  DB authorization rows and the filesystem cannot be locked in one
        # atomic transaction, so every ready/download path remains live-authorized too.
        _mark_artifact_ready(artifact_id, store.inspect(storage_key))
    except (
        ArtifactStoreUnavailable,
        AuthorizationUnavailable,
        ObjectPublicationUncertain,
    ) as exc:
        # If bytes are already durable, authorization/DB uncertainty is not a proven
        # denial. Keep validating as a non-downloadable marker; only the normal
        # publisher path may ever supply the missing completion receipt.
        raise ArtifactStoreUnavailable("文件发布状态待协调，请稍后重试") from exc
    except Exception as exc:  # noqa: BLE001 - internal detail is deliberately hidden
        # Never delete a formally published object here.  A known FileError is a
        # deterministic rejection and may become failed; an unexpected exception after
        # publication is UNKNOWN and must remain validating for reconciliation.
        if isinstance(exc, FileError) or not object_published:
            try:
                _mark_artifact_failed(artifact_id)
            except Exception:  # noqa: BLE001 - preserve the stable public error on DB outage
                pass
        raise FileError("文件发布失败，请稍后重试") from exc
    return artifact_info(artifact_id)


@dataclass
class _ProvenanceAuthorizationState:
    memo: dict[str, bool]
    verified_sources: dict[str, dict[str, Any]] = field(default_factory=dict)
    nodes: int = 0
    work: int = 0
    exhausted: bool = False


def _memoize_access(
    state: _ProvenanceAuthorizationState,
    file_id: str,
    result: bool,
) -> bool:
    if file_id not in state.memo:
        if len(state.memo) >= _MAX_PROVENANCE_AUTH_MEMO:
            state.exhausted = True
            return False
        state.memo[file_id] = result
    return result


def access_allowed(
    file_id: str,
    user_ctx: Any,
    *,
    _state: _ProvenanceAuthorizationState | None = None,
    _path: frozenset[str] | None = None,
    _depth: int = 0,
) -> bool:
    """Re-authorize one provenance DAG under global work and cycle budgets."""
    try:
        subject = stable_owner_sub(user_ctx)
    except FileError:
        return False
    fid = _check_id(file_id)
    state = (
        _state
        if isinstance(_state, _ProvenanceAuthorizationState)
        else _ProvenanceAuthorizationState(memo={})
    )
    path = frozenset(_path or ())
    if state.exhausted:
        return False
    # A node currently on this recursion path is a real cycle, even if another
    # completed branch has a memoized result for the same Artifact.
    if fid in path:
        return False
    if fid in state.memo:
        return state.memo[fid]
    if (
        _depth > _MAX_PROVENANCE_DERIVATION_DEPTH
        or state.nodes >= _MAX_PROVENANCE_AUTH_NODES
        or len(path) >= _MAX_PROVENANCE_AUTH_PATH
    ):
        state.exhausted = True
        return False
    state.nodes += 1
    current_path = path | {fid}
    try:
        meta = _find_artifact_meta(fid, require_ready=False)
    except ArtifactUnavailable:
        # A malformed or no-longer-verifiable binding is a deterministic deny at
        # this boolean ACL boundary.  Infrastructure exceptions still propagate
        # so callers that distinguish UNKNOWN from DENIED can fail separately.
        return _memoize_access(state, fid, False)
    scope = (meta or {}).get("access_scope") or {}
    snapshots = scope.get("source_access_snapshots")
    state.work += 1 + (len(snapshots) if isinstance(snapshots, list) else 0)
    if state.work > _MAX_PROVENANCE_AUTH_WORK:
        state.exhausted = True
        return False
    result = _evaluate_artifact_access(
        fid=fid,
        meta=meta,
        user_ctx=user_ctx,
        subject=subject,
        state=state,
        path=current_path,
        depth=_depth,
    )
    return _memoize_access(state, fid, result)


def _authorize_source_snapshot(
    payload: dict[str, Any],
    *,
    user_ctx: Any,
    state: _ProvenanceAuthorizationState,
    path: frozenset[str],
    depth: int,
) -> bool:
    """Re-read and authorize one signed Artifact source under a shared DAG budget."""
    source_id = payload.get("source_artifact_id")
    if not isinstance(source_id, str) or state.exhausted:
        return False
    source_meta = state.verified_sources.get(source_id)
    if source_meta is None:
        if len(state.verified_sources) >= _MAX_PROVENANCE_AUTH_MEMO:
            state.exhausted = True
            return False
        try:
            source_meta = _artifact_meta(source_id, require_ready=True)
            _verify_artifact(source_meta)
        except (ArtifactStoreUnavailable, AuthorizationUnavailable):
            raise
        except ArtifactUnavailable as exc:
            if exc.reason_code in {"state_audit_failed", "binding_resign_failed"}:
                raise AuthorizationUnavailable(
                    "来源 Artifact 状态暂时无法确认"
                ) from exc
            return False
        except FileError:
            return False
        state.verified_sources[source_id] = source_meta
    return (
        source_meta.get("file_id") == source_id
        and source_meta.get("status") == "ready"
        and source_meta.get("sha256") == payload.get("source_sha256")
        and source_meta.get("owner_sub") == payload.get("owner_sub")
        and artifact_provenance.artifact_snapshot_matches_scope(
            payload,
            source_meta.get("access_scope") or {},
            source_sha256=source_meta.get("sha256") or "",
        )
        and access_allowed(
            source_id,
            user_ctx,
            _state=state,
            _path=path,
            _depth=depth,
        )
    )


def _reauthorize_provenance_scope(
    owner: VerifiedArtifactOwner,
    scope: dict[str, Any],
) -> list[str]:
    """Reload the principal and every Artifact source before read/publish checkpoints."""
    try:
        current = _verified_owner_context(owner)
        owner_sub = stable_owner_sub(current)
        state = _ProvenanceAuthorizationState(memo={})

        def authorize_source(payload: dict[str, Any]) -> bool:
            return _authorize_source_snapshot(
                payload,
                user_ctx=current,
                state=state,
                path=frozenset(),
                depth=0,
            )

        allowed = artifact_provenance.current_scope_covers(
            scope,
            current,
            source_artifact_authorizer=authorize_source,
        )
        source_ids = artifact_provenance.source_artifact_ids(scope)
        if (
            not allowed
            or state.exhausted
            or any(
                (meta := state.verified_sources.get(source_id)) is None
                or meta.get("owner_sub") != owner_sub
                for source_id in source_ids
            )
        ):
            raise ProvenanceRequired("生成文件当前来源授权已失效，已停止处理")
        return source_ids
    except ProvenanceRequired:
        raise
    except (ArtifactStoreUnavailable, AuthorizationUnavailable) as exc:
        raise AuthorizationUnavailable(
            "生成文件当前授权暂时无法确认，已停止处理"
        ) from exc
    except (FileError, artifact_provenance.ProvenanceError) as exc:
        raise ProvenanceRequired("生成文件当前来源授权已失效，已停止处理") from exc
    except Exception as exc:  # noqa: BLE001 - live path fails closed on uncertainty
        raise AuthorizationUnavailable("生成文件当前授权暂时无法确认，已停止处理") from exc


def _reconcile_ready_authorized(row: AgentArtifact) -> bool:
    """Apply the publisher's live principal/scope contract before crash recovery."""
    try:
        _require_artifact_binding(_binding_metadata_from_row(row))
    except FileError:
        return False
    extra = row.extra_meta if isinstance(row.extra_meta, dict) else {}
    user_id = extra.get("_publisher_user_id")
    token_version = extra.get("_publisher_token_version")
    if (
        isinstance(user_id, bool)
        or not isinstance(user_id, int)
        or isinstance(token_version, bool)
        or not isinstance(token_version, int)
    ):
        return False
    try:
        with SessionLocal() as db:
            user = db.scalar(select(SysUser).where(
                SysUser.id == user_id,
                SysUser.username == row.owner_sub,
            ))
            if user is None:
                return False
            owner = VerifiedArtifactOwner(user.id, row.owner_sub, token_version)
        _verified_owner_sub(owner)
        if row.kind == "upload":
            return not list(row.source_ids or [])
        if row.kind != "generated":
            return False
        if not get_settings().agent_artifact_v2_enabled:
            raise AuthorizationUnavailable(
                "Artifact v2 暂停期间无法确认生成制品来源授权"
            )
        source_ids = _reauthorize_provenance_scope(
            owner,
            dict(row.access_scope or {}),
        )
        return source_ids == list(row.source_ids or [])
    except AuthorizationUnavailable:
        raise
    except FileError:
        return False


def _evaluate_artifact_access(
    *,
    fid: str,
    meta: dict[str, Any] | None,
    user_ctx: Any,
    subject: str,
    state: _ProvenanceAuthorizationState,
    path: frozenset[str],
    depth: int,
) -> bool:
    from app import config, permissions

    if meta is None and _is_legacy_id(fid):
        # A legacy sidecar is caller-writable storage metadata, not an authenticated
        # binding to today's SysUser row.  Even an exact same-name account must not
        # silently adopt it; adoption/break-glass is a separate operator workflow.
        return False
    if meta is None:
        raise FileError("文件不存在或已清理")
    owner_ok = meta.get("operated_by") == subject
    if not owner_ok:
        return False
    scope = meta.get("access_scope") or {}
    policy = scope.get("policy")
    if meta.get("kind") == "upload":
        return policy == "owner_only"
    if policy == "unclassified_deny":
        return False
    if policy == "provenance_guarded":
        def authorize_source(payload: dict[str, Any]) -> bool:
            return _authorize_source_snapshot(
                payload,
                user_ctx=user_ctx,
                state=state,
                path=path,
                depth=depth + 1,
            )

        return artifact_provenance.current_scope_covers(
            scope,
            user_ctx,
            source_artifact_authorizer=authorize_source,
        )
    if policy != "current_scope_dominates" or scope.get("version") != 1:
        return False
    if not config.ENABLE_RBAC and user_ctx.role == config.PHASE1_BYPASS_ROLE:
        return True
    permission_source = (
        permissions.effective(user_ctx.role, None)
        if user_ctx.permissions is None
        else user_ctx.permissions
    )
    current = permissions.runtime_safe(permission_source)
    required = scope.get("required_permissions")
    if not isinstance(required, list) or not all(isinstance(key, str) for key in required):
        return False
    allowed_keys = set(permissions.DATA_GROUPS) | set(permissions.PAGE_KEYS)
    if any(key not in allowed_keys or not current.get(key, False) for key in required):
        return False
    data_snapshot = scope.get("data_permissions")
    page_snapshot = scope.get("page_permissions")
    visible_snapshot = scope.get("visible_field_groups")
    if not isinstance(data_snapshot, dict) or set(data_snapshot) != set(permissions.DATA_GROUPS):
        return False
    if not isinstance(page_snapshot, dict) or set(page_snapshot) != set(permissions.PAGE_KEYS):
        return False
    if not isinstance(visible_snapshot, list) or not all(isinstance(group, str) for group in visible_snapshot):
        return False
    if any(bool(granted) and not current.get(key, False) for key, granted in data_snapshot.items()):
        return False
    if any(bool(granted) and not current.get(key, False) for key, granted in page_snapshot.items()):
        return False
    current_visible_groups = {
        group
        for key, groups in permissions.DATA_GROUPS.items()
        if current.get(key, False)
        for group in groups
    }
    if not set(visible_snapshot).issubset(current_visible_groups):
        return False
    snapshot_row_scope = scope.get("row_scope")
    if snapshot_row_scope not in {"all", "own_customers"}:
        return False
    current_is_own = bool(current.get("own_customers_only", False))
    # An all-customers result cannot be reopened after the account becomes own-only.
    if snapshot_row_scope == "all" and current_is_own:
        return False
    if snapshot_row_scope == "own_customers" and current_is_own:
        snapshot_subject = scope.get("row_subject")
        current_subject = _canonical_salesperson_subject(user_ctx.salesperson_name)
        if (
            not isinstance(snapshot_subject, str)
            or not snapshot_subject
            or not current_subject
            or snapshot_subject != current_subject
        ):
            return False
    return True


def _authorized_owner_id(file_id: str, owner: VerifiedArtifactOwner) -> str:
    """Combine live-principal validation, owner ACL and current-scope recheck."""
    fid = _check_id(file_id)
    ctx = _verified_owner_context(owner)
    try:
        allowed = access_allowed(fid, ctx)
    except ArtifactV2Disabled:
        raise
    except (ArtifactStoreUnavailable, AuthorizationUnavailable):
        raise
    except FileError as exc:
        raise ArtifactUnavailable(
            "文件不存在或无权访问", "not_found_or_forbidden"
        ) from exc
    if not allowed:
        raise ArtifactUnavailable("文件不存在或无权访问", "not_found_or_forbidden")
    return fid


def _cell_str(v) -> str:
    if v is None:
        return ""
    s = v.strftime("%Y-%m-%d") if hasattr(v, "strftime") else str(v)
    return s[:_CELL_TRUNC]


def _safe_spreadsheet_value(value):
    """Neutralize formula-like untrusted text while leaving real numbers unchanged."""
    if isinstance(value, str) and value.lstrip().startswith(_FORMULA_PREFIXES):
        return f"'{value}"
    return value


def _canonical_salesperson_subject(value: Any) -> str | None:
    """Canonical identity bound to own-customer row scope (not a display label)."""
    if value is None:
        return None
    normalized = unicodedata.normalize("NFKC", str(value))
    normalized = re.sub(r"\s+", " ", normalized).strip().casefold()
    return normalized or None


def save_upload(content: bytes, filename: str, owner: VerifiedArtifactOwner) -> dict:
    """保存上传文件（多格式），返回 file_id + 类型/概览（供注入对话上下文）。"""
    require_artifact_v2_enabled()
    _verified_owner_sub(owner)
    safe_name = _safe_filename(filename, "上传文件")
    ext = _ext_of(safe_name)
    if ext == "xls":
        raise FileError("不支持旧版 .xls，请用 Excel 另存为 .xlsx")
    if ext not in _ALLOWED_EXT:
        raise FileError(f"不支持的文件类型 .{ext}（支持：Excel/Word/PDF/txt/图片）")
    if len(content) > _MAX_UPLOAD_MB * 1024 * 1024:
        raise FileError(f"文件超过 {_MAX_UPLOAD_MB}MB 上限")
    upload_scope = artifact_provenance.classify_upload_access_scope(content, ext)

    extra_meta: dict = {}
    if ext == "xlsx":
        # xlsx 校验可解析并带回 sheet 概览
        try:
            wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
            extra_meta["sheets"] = [{"name": ws.title, "n_rows": ws.max_row or 0,
                                     "n_cols": ws.max_column or 0} for ws in wb.worksheets]
            wb.close()
        except Exception as exc:  # noqa: BLE001
            raise FileError(f"无法解析 xlsx（可能损坏，请用 Excel 重新另存）: {type(exc).__name__}") from exc

    artifact = _publish_artifact(
        content,
        safe_name,
        kind="upload",
        owner=owner,
        provenance_scope=upload_scope,
        extra_meta=extra_meta,
    )
    out = {"file_id": artifact["file_id"], "filename": safe_name, "ext": ext,
           "file_kind": ("表格" if ext == "xlsx" else "Word" if ext == "docx"
                         else "PDF" if ext == "pdf" else "图片" if ext in _IMG_EXT else "文本"),
           "artifact": _artifact_ref(artifact)}
    if "sheets" in extra_meta:
        out["sheets"] = extra_meta["sheets"]
    return out


# ============================================================
# 读取（xlsx 结构化 + 通用 read_document）
# ============================================================

def _require_xlsx(fid: str, meta: dict):
    if meta.get("ext") != "xlsx":
        raise FileError(f"该文件是 .{meta.get('ext')}，结构化读取仅限 Excel；请用 read_document 读取内容")


def inspect_file(file_id: str, owner: VerifiedArtifactOwner) -> dict:
    """看 Excel 结构：sheet 列表 + 每 sheet 前几行原样预览。"""
    fid = _authorized_owner_id(file_id, owner)
    meta, stored = _load_verified_artifact(fid)
    _require_xlsx(fid, meta)
    wb = load_workbook(BytesIO(stored.content), read_only=True, data_only=True)
    sheets = []
    for ws in wb.worksheets[:5]:
        preview = [[_cell_str(c.value) for c in row]
                   for row in ws.iter_rows(min_row=1, max_row=_PREVIEW_ROWS,
                                           max_col=min(ws.max_column or 1, _PREVIEW_COLS))]
        sheets.append({"name": ws.title, "n_rows": ws.max_row or 0,
                       "n_cols": ws.max_column or 0, "preview_rows_1_to_n": preview})
    wb.close()
    return {"file_id": fid, "filename": meta.get("filename"), "sheets": sheets,
            "note": "preview 为前几行原样数据(1-based 行号)；更多行用 read_file_rows"}


def read_rows(file_id: str, sheet: str | None, start_row: int, max_rows: int,
              owner: VerifiedArtifactOwner) -> dict:
    """分页读取 Excel 行（1-based）。"""
    fid = _authorized_owner_id(file_id, owner)
    meta, stored = _load_verified_artifact(fid)
    _require_xlsx(fid, meta)
    wb = load_workbook(BytesIO(stored.content), read_only=True, data_only=True)
    try:
        ws = wb[sheet] if sheet else wb.worksheets[0]
    except KeyError:
        names = [w.title for w in wb.worksheets]
        wb.close()
        raise FileError(f"sheet 不存在: {sheet!r}，可选: {names}")
    start = max(int(start_row or 1), 1)
    n = min(int(max_rows or 50), _MAX_READ_ROWS)
    rows = [[_cell_str(c.value) for c in row]
            for row in ws.iter_rows(min_row=start, max_row=start + n - 1,
                                    max_col=min(ws.max_column or 1, 30))]
    total = ws.max_row or 0
    wb.close()
    return {"file_id": fid, "sheet": ws.title, "start_row": start,
            "rows": rows, "total_rows": total}


def preview(file_id: str, owner: VerifiedArtifactOwner, max_rows: int = 200) -> dict:
    """文件预览（前端在线预览用）：xlsx 返回各 sheet 行数据（截断 max_rows×30列）；
    图片返回 kind=image（前端走下载端点取图）；其余 kind=other（仅可下载）。
    归属与当前权限在服务层统一校验。"""
    fid = _authorized_owner_id(file_id, owner)
    meta, stored = _load_verified_artifact(fid)
    ext = meta.get("ext", "")
    filename = meta.get("filename", f"{fid}.{ext}")
    if ext != "xlsx":
        kind = "image" if ext in _IMG_EXT else "other"
        return {"file_id": fid, "filename": filename, "kind": kind, "ext": ext}
    # 坏/半损 xlsx 可能通过上传校验(只读维度)却在逐格迭代时抛 ParseError/BadZipFile(非 FileError)，
    # 不裹会让预览端点裸冒 500 → 统一转 FileError，端点据此返干净 404（与 save_upload 一致）
    try:
        wb = load_workbook(
            BytesIO(stored.content),
            read_only=True,
            data_only=True,
        )
        sheets = []
        for ws in wb.worksheets[:10]:
            total = ws.max_row or 0
            rows = [[_cell_str(c.value) for c in row]
                    for row in ws.iter_rows(min_row=1, max_row=min(total, max_rows),
                                            max_col=min(ws.max_column or 1, 30))]
            sheets.append({"name": ws.title, "rows": rows,
                           "total_rows": total, "truncated": total > max_rows})
        wb.close()
    except FileError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise FileError(f"无法解析 xlsx（可能损坏，请重新另存）: {type(exc).__name__}") from exc
    return {"file_id": fid, "filename": filename, "kind": "table", "ext": ext, "sheets": sheets}


def _read_docx(source: Any) -> str:
    from docx import Document
    doc = Document(source)
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for ti, table in enumerate(doc.tables):
        parts.append(f"[表格{ti + 1}]")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_pdf(source: Any) -> tuple[str, bool]:
    """返回 (文本, 是否疑似扫描件)。文字层为空/极少 → 扫描件，转视觉。"""
    import pdfplumber
    parts: list[str] = []
    with pdfplumber.open(source) as pdf:
        for pi, page in enumerate(pdf.pages[:30]):
            txt = page.extract_text() or ""
            parts.append(f"[第{pi + 1}页]\n{txt}" if txt.strip() else f"[第{pi + 1}页](无文字层)")
            for table in page.extract_tables() or []:
                for row in table:
                    cells = [(c or "").strip() for c in row]
                    if any(cells):
                        parts.append(" | ".join(cells))
    text = "\n".join(parts)
    scanned = len(re.sub(r"\s", "", text.replace("无文字层", ""))) < 20
    return text, scanned


def _read_image_or_scanned(path: Path, hint: str) -> str:
    """图片/扫描件 → Qwen-VL 识别（无 key 优雅降级）。"""
    from app.agent import provider
    try:
        return provider.vision_extract([path], hint)
    except provider.VisionNotConfigured:
        return ("【未配置视觉模型】该文件是图片/扫描件，需配置 VISION_API_KEY（通义 Qwen-VL）"
                "后才能识别。文字版 Word/Excel/PDF/txt 不受影响。")


def _read_image_or_scanned_bytes(content: bytes, ext: str, hint: str) -> str:
    """Bridge immutable verified bytes to the pathname-only Vision adapter."""
    if ext not in _IMG_EXT | {"pdf"}:
        raise FileError("视觉文件类型无效")
    with tempfile.TemporaryDirectory(prefix="agent-artifact-vision-") as directory:
        path = Path(directory) / f"verified.{ext}"
        flags = os.O_WRONLY | os.O_CREAT | os.O_EXCL | getattr(os, "O_CLOEXEC", 0)
        descriptor = os.open(path, flags, 0o600)
        try:
            with os.fdopen(descriptor, "wb", closefd=True) as handle:
                descriptor = -1
                handle.write(content)
        finally:
            if descriptor >= 0:
                os.close(descriptor)
        return _read_image_or_scanned(path, hint)


def read_document(file_id: str, owner: VerifiedArtifactOwner) -> dict:
    """通用读取：把任意支持格式抽成文本喂给模型（拆件/解析由模型完成）。"""
    fid = _authorized_owner_id(file_id, owner)
    meta, stored = _load_verified_artifact(fid)
    ext = meta.get("ext", "")
    vision_used = False
    if ext in _TEXT_EXT:
        text = stored.content.decode("utf-8", errors="replace")
    elif ext == "xlsx":
        wb = load_workbook(
            BytesIO(stored.content),
            read_only=True,
            data_only=True,
        )
        chunks = []
        for ws in wb.worksheets:
            chunks.append(f"[工作表 {ws.title}]")
            for row in ws.iter_rows(max_row=min(ws.max_row or 0, 300), max_col=min(ws.max_column or 1, 30)):
                cells = [_cell_str(c.value) for c in row]
                if any(cells):
                    chunks.append(" | ".join(cells))
        wb.close()
        text = "\n".join(chunks)
    elif ext == "docx":
        text = _read_docx(BytesIO(stored.content))
    elif ext == "pdf":
        text, scanned = _read_pdf(BytesIO(stored.content))
        if scanned:
            vision_used = True
            text = _read_image_or_scanned_bytes(
                stored.content,
                "pdf",
                "这是一份扫描件，请逐字识别其中的全部文本、表格、型号与参数。",
            )
    elif ext in _IMG_EXT:
        vision_used = True
        text = _read_image_or_scanned_bytes(
            stored.content,
            ext,
            "请识别图片中的全部文字、表格、设备型号、品牌与参数配置，按原结构输出。",
        )
    else:
        raise FileError(f"不支持读取 .{ext}")

    truncated = len(text) > _DOC_CHAR_CAP
    return {"file_id": fid, "filename": meta.get("filename"), "ext": ext,
            "vision_used": vision_used, "truncated": truncated,
            "content": text[:_DOC_CHAR_CAP]}


# ============================================================
# 写（模板回填 write_excel / 美化报表 write_report）
# ============================================================

def _col_index(col) -> int:
    """列定位：支持 "A"/"G" 字母或 1-based 数字。"""
    if isinstance(col, int) or (isinstance(col, str) and col.isdigit()):
        idx = int(col)
        if idx < 1 or idx > 16384:
            raise FileError(f"列号超界: {col}")
        return idx
    try:
        return column_index_from_string(str(col).strip().upper())
    except Exception as exc:  # noqa: BLE001
        raise FileError(f"无法识别列: {col!r}（用字母如 'G' 或 1-based 数字）") from exc


def _bounded_render_value_bytes(value: Any) -> int:
    """Validate one renderer scalar and return its bounded textual byte cost."""
    if value is None or isinstance(value, bool):
        return 0
    if isinstance(value, int):
        return len(str(value))
    if isinstance(value, float):
        if not math.isfinite(value):
            raise FileError("生成文件不支持 NaN 或无穷数值")
        return len(repr(value))
    if isinstance(value, Decimal):
        if not value.is_finite():
            raise FileError("生成文件不支持 NaN 或无穷数值")
        return len(str(value))
    if isinstance(value, (date, datetime)):
        return len(value.isoformat().encode("utf-8"))
    if isinstance(value, str):
        size = len(value.encode("utf-8"))
        if size > _MAX_RENDER_VALUE_BYTES:
            raise FileError("生成文件单个文本值超过大小预算")
        return size
    raise FileError(f"生成文件不支持值类型 {type(value).__name__}")


def _bounded_optional_text(value: Any, label: str, *, max_bytes: int) -> int:
    if value is None:
        return 0
    if not isinstance(value, str):
        raise FileError(f"{label} 必须是字符串")
    size = len(value.encode("utf-8"))
    if size > max_bytes:
        raise FileError(f"{label} 超过大小预算")
    return size


def _validate_excel_write_shape(
    *,
    sheet: str | None,
    cells: list[dict],
    output_name: str | None,
) -> None:
    total_bytes = _bounded_optional_text(
        output_name, "output_name", max_bytes=1024
    )
    if sheet is not None:
        total_bytes += _bounded_optional_text(sheet, "sheet", max_bytes=128)
        if not sheet or len(sheet) > 31 or re.search(r"[\\/*?:\[\]]", sheet):
            raise FileError("sheet 不是有效的 Excel 工作表名称")
    if not isinstance(cells, list) or not cells:
        raise FileError("cells 必须是非空数组")
    if len(cells) > _MAX_WRITE_CELLS:
        raise FileError(f"单次最多写 {_MAX_WRITE_CELLS} 个单元格")
    for item in cells:
        if not isinstance(item, dict) or set(item) != {"row", "col", "value"}:
            raise FileError("cells 项格式错（只允许 row/col/value）")
        row = item["row"]
        if isinstance(row, bool) or not isinstance(row, int):
            raise FileError("cells.row 必须是整数")
        if row < 1 or row > 1_048_576:
            raise FileError(f"行号超界: {row}")
        if isinstance(item["col"], bool):
            raise FileError("cells.col 必须是列字母或整数")
        _col_index(item["col"])
        total_bytes += _bounded_render_value_bytes(item["value"])
        if total_bytes > _MAX_RENDER_TEXT_BYTES:
            raise FileError("生成文件文本总量超过预算")


def _validate_report_shape(
    *,
    title: str | None,
    headers: list[str],
    rows: list[list],
    output_name: str | None,
    money_cols: list[int] | None,
) -> None:
    total_bytes = _bounded_optional_text(title, "title", max_bytes=4096)
    total_bytes += _bounded_optional_text(
        output_name, "output_name", max_bytes=1024
    )
    if not isinstance(headers, list) or not headers:
        raise FileError("headers 必须是非空数组")
    if len(headers) > _MAX_REPORT_COLUMNS:
        raise FileError(f"报表最多 {_MAX_REPORT_COLUMNS} 列")
    for header in headers:
        if not isinstance(header, str):
            raise FileError("headers 每项必须是字符串")
        total_bytes += _bounded_render_value_bytes(header)
        if total_bytes > _MAX_RENDER_TEXT_BYTES:
            raise FileError("生成文件文本总量超过预算")
    if not isinstance(rows, list):
        raise FileError("rows 必须是二维数组")
    if len(rows) > _MAX_REPORT_ROWS:
        raise FileError(f"报表最多 {_MAX_REPORT_ROWS} 行")
    if len(rows) * len(headers) > _MAX_REPORT_CELLS:
        raise FileError("报表单元格总量超过预算")
    for row in rows:
        if not isinstance(row, list) or len(row) > len(headers):
            raise FileError("rows 每行必须是数组且不能宽于 headers")
        for value in row:
            total_bytes += _bounded_render_value_bytes(value)
            if total_bytes > _MAX_RENDER_TEXT_BYTES:
                raise FileError("生成文件文本总量超过预算")
    if money_cols is not None:
        if (
            not isinstance(money_cols, list)
            or len(money_cols) > _MAX_MONEY_COLUMNS
            or any(
                isinstance(index, bool)
                or not isinstance(index, int)
                or index < 0
                or index >= len(headers)
                for index in money_cols
            )
            or money_cols != sorted(set(money_cols))
        ):
            raise FileError("money_cols 必须是升序、不重复且位于 headers 范围内")


def write_excel(base_file_id: str | None, sheet: str | None,
                cells: list[dict], output_name: str | None,
                owner: VerifiedArtifactOwner, *,
                provenance: Any | None = None) -> dict:
    """按模型指令写单元格，产出新文件（不动原件）。用于**回填客户模板**（保留原格式）。"""
    require_artifact_v2_enabled()
    owner_sub = _verified_owner_sub(owner)
    # Authorize the source before validating model-supplied edit instructions.  This
    # keeps non-owner/unknown source IDs indistinguishable at the service boundary.
    base = _canonical_source_id(base_file_id, owner) if base_file_id else None
    if provenance is None:
        raise ProvenanceRequired("生成文件缺少可验证的数据来源，已停止发布")
    _validate_excel_write_shape(
        sheet=sheet,
        cells=cells,
        output_name=output_name,
    )
    try:
        provenance_scope = artifact_provenance.consume_evidence(
            provenance,
            owner_sub=owner_sub,
            expected_fingerprint=artifact_provenance.excel_edit_fingerprint(
                base_file_id=base,
                sheet=sheet,
                cells=cells,
                output_name=output_name,
            ),
        )
        proven_source_ids = artifact_provenance.source_artifact_ids(provenance_scope)
    except artifact_provenance.ProvenanceError as exc:
        raise ProvenanceRequired("生成文件来源证明无效，已停止发布") from exc
    if base is not None and base not in proven_source_ids:
        raise ProvenanceRequired("模板来源未包含在生成证明中，已停止发布")
    if provenance_scope.get("classification") == "identity_only":
        raise ProvenanceRequired("identity_only 模板不能证明模型写入内容，已停止发布")
    # This checkpoint is deliberately after evidence consumption and before the
    # workbook is opened.  Permission, row-scope, status, owner, hash or nested scope
    # drift therefore cannot expose source bytes to the renderer.
    _reauthorize_provenance_scope(owner, provenance_scope)
    base_name = ""
    if base:
        bmeta, stored = _load_verified_artifact(base)
        _require_xlsx(base, bmeta)
        # Preserve styles/formulas from the exact bytes verified on one no-follow
        # handle; never reopen the mutable object path after validation.
        wb = load_workbook(BytesIO(stored.content))
        base_name = bmeta.get("filename", "")
    else:
        wb = Workbook()

    ws = (wb[sheet] if sheet in wb.sheetnames else wb.create_sheet(sheet)) if sheet else wb.worksheets[0]

    written = 0
    for c in cells:
        row, col = c["row"], _col_index(c["col"])
        ws.cell(row=row, column=col, value=_safe_spreadsheet_value(c.get("value")))
        written += 1

    name = output_name or (f"回填_{base_name}" if base_name else "结果.xlsx")
    if not name.lower().endswith(".xlsx"):
        name += ".xlsx"
    buffer = BytesIO()
    wb.save(buffer)
    wb.close()
    artifact = _publish_artifact(
        buffer.getvalue(),
        name,
        kind="generated",
        owner=owner,
        provenance_scope=provenance_scope,
        extra_meta={"base_file_id": base_file_id},
    )
    ref = _artifact_ref(artifact)
    download_url = f"/api/agent/files/{artifact['file_id']}"
    return {"file_id": artifact["file_id"], "filename": artifact["filename"],
            "cells_written": written, "download_url": download_url, "artifact": ref}


# 美化报表样式常量
_HEADER_FILL = PatternFill("solid", fgColor="4F46E5")
_HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
_TITLE_FONT = Font(bold=True, size=14, color="111827")
_ZEBRA_FILL = PatternFill("solid", fgColor="F5F6FF")
_WARN_FILL = PatternFill("solid", fgColor="FFF3E0")    # 需确认=橙
_BAD_FILL = PatternFill("solid", fgColor="FDECEA")     # 未找到=红
_THIN = Side(style="thin", color="E5E7EB")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
_WARN_KW = ("需确认", "待确认", "确认")
_BAD_KW = ("未找到", "无库存", "不存在", "未匹配")


def write_report(title: str | None, headers: list[str], rows: list[list],
                 output_name: str | None, owner: VerifiedArtifactOwner,
                 money_cols: list[int] | None = None, *,
                 provenance: Any | None = None) -> dict:
    """生成**美化报表**（BOM/报价单等）：表头配色、边框、自适应列宽、金额格式、
    冻结表头、斑马纹、状态行条件配色。headers=列名；rows=与列对齐的二维数组；
    money_cols=金额列的 0-based 下标（数字格式+右对齐）。"""
    require_artifact_v2_enabled()
    owner_sub = _verified_owner_sub(owner)
    if provenance is None:
        raise ProvenanceRequired("生成文件缺少可验证的数据来源，已停止发布")
    _validate_report_shape(
        title=title,
        headers=headers,
        rows=rows,
        output_name=output_name,
        money_cols=money_cols,
    )
    expected_fingerprint = artifact_provenance.report_fingerprint(
        title=title,
        headers=headers,
        rows=rows,
        output_name=output_name,
        money_cols=money_cols,
    )
    try:
        provenance_scope = artifact_provenance.consume_evidence(
            provenance,
            owner_sub=owner_sub,
            expected_fingerprint=expected_fingerprint,
        )
    except artifact_provenance.ProvenanceError as exc:
        raise ProvenanceRequired("生成文件来源证明无效，已停止发布") from exc
    if provenance_scope.get("classification") != "business_content":
        raise ProvenanceRequired("来源只能证明模板身份，不能证明报表业务内容")
    money = set(money_cols or [])
    ncol = len(headers)

    wb = Workbook()
    ws = wb.worksheets[0]
    ws.title = "报表"
    r0 = 1
    if title:
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncol)
        tc = ws.cell(row=1, column=1, value=_safe_spreadsheet_value(title))
        tc.font = _TITLE_FONT
        tc.alignment = Alignment(horizontal="left", vertical="center")
        ws.row_dimensions[1].height = 26
        r0 = 2

    # 表头
    for j, h in enumerate(headers, start=1):
        c = ws.cell(row=r0, column=j, value=_safe_spreadsheet_value(str(h)))
        c.fill, c.font, c.border = _HEADER_FILL, _HEADER_FONT, _BORDER
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[r0].height = 22

    widths = [len(str(h)) for h in headers]
    for i, row in enumerate(rows):
        rr = r0 + 1 + i
        rowtext = " ".join(str(v) for v in row if v is not None)
        hit_bad = any(k in rowtext for k in _BAD_KW)
        hit_warn = not hit_bad and any(k in rowtext for k in _WARN_KW)
        base_fill = _BAD_FILL if hit_bad else _WARN_FILL if hit_warn else (_ZEBRA_FILL if i % 2 else None)
        for j in range(ncol):
            val = row[j] if j < len(row) else None
            c = ws.cell(row=rr, column=j + 1, value=_safe_spreadsheet_value(val))
            c.border = _BORDER
            if base_fill:
                c.fill = base_fill
            if j in money and isinstance(val, (int, float)):
                c.number_format = '#,##0.00'
                c.alignment = Alignment(horizontal="right")
            else:
                c.alignment = Alignment(vertical="center", wrap_text=isinstance(val, str) and len(str(val)) > 20)
            widths[j] = max(widths[j], min(len(_cell_str(val)), 50))

    for j, w in enumerate(widths, start=1):
        # 中文偏宽：粗略 ×1.6 + 余量；上限防超宽
        ws.column_dimensions[get_column_letter(j)].width = min(max(w * 1.6 + 2, 8), 60)
    ws.freeze_panes = ws.cell(row=r0 + 1, column=1)

    name = output_name or (title or "报表")
    if not name.lower().endswith(".xlsx"):
        name += ".xlsx"
    buffer = BytesIO()
    wb.save(buffer)
    wb.close()
    artifact = _publish_artifact(
        buffer.getvalue(),
        name,
        kind="generated",
        owner=owner,
        provenance_scope=provenance_scope,
        extra_meta={"report": True},
    )
    ref = _artifact_ref(artifact)
    download_url = f"/api/agent/files/{artifact['file_id']}"
    return {"file_id": artifact["file_id"], "filename": artifact["filename"],
            "rows_written": len(rows), "download_url": download_url, "artifact": ref}


def _mint_report_provenance(
    owner: VerifiedArtifactOwner,
    *,
    title: str | None,
    headers: list[str],
    rows: list[list],
    output_name: str | None,
    money_cols: list[int] | None,
    contained_resources: set[str],
    contained_fields: set[str],
    required_positive_keys: set[str],
) -> artifact_provenance.TrustedEvidence:
    """Trusted Query Broker seam; deliberately absent from tool and HTTP schemas."""
    _validate_report_shape(
        title=title,
        headers=headers,
        rows=rows,
        output_name=output_name,
        money_cols=money_cols,
    )
    ctx = _verified_owner_context(owner)
    owner_sub = stable_owner_sub(ctx)
    return artifact_provenance.mint_server_evidence(
        owner_sub=owner_sub,
        content_fingerprint_value=artifact_provenance.report_fingerprint(
            title=title,
            headers=headers,
            rows=rows,
            output_name=output_name,
            money_cols=money_cols,
        ),
        contained_resources=contained_resources,
        contained_fields=contained_fields,
        required_positive_keys=required_positive_keys,
        row_subject=_canonical_salesperson_subject(ctx.salesperson_name),
        own_customers_only=bool((ctx.permissions or {}).get("own_customers_only")),
    )


def _mint_report_from_artifacts(
    owner: VerifiedArtifactOwner,
    *,
    source_ids: list[str],
    title: str | None,
    headers: list[str],
    rows: list[list],
    output_name: str | None,
    money_cols: list[int] | None,
) -> artifact_provenance.TrustedEvidence:
    """Trusted deterministic-transform seam; never exposed to model tool arguments."""
    _validate_report_shape(
        title=title,
        headers=headers,
        rows=rows,
        output_name=output_name,
        money_cols=money_cols,
    )
    requested_ids = artifact_provenance.validate_source_artifact_ids(source_ids)
    ctx = _verified_owner_context(owner)
    owner_sub = stable_owner_sub(ctx)
    canonical_ids = [_canonical_source_id(source_id, owner) for source_id in requested_ids]
    source_metas = [_load_meta(source_id) for source_id in canonical_ids]
    return artifact_provenance.mint_artifact_evidence(
        owner_sub=owner_sub,
        content_fingerprint_value=artifact_provenance.report_fingerprint(
            title=title,
            headers=headers,
            rows=rows,
            output_name=output_name,
            money_cols=money_cols,
        ),
        source_metas=source_metas,
    )


def _mint_excel_from_artifacts(
    owner: VerifiedArtifactOwner,
    *,
    source_ids: list[str],
    base_file_id: str | None,
    sheet: str | None,
    cells: list[dict],
    output_name: str | None,
) -> artifact_provenance.TrustedEvidence:
    """Trusted Change Plan seam; never exposed to model or request fields."""
    requested_ids = artifact_provenance.validate_source_artifact_ids(source_ids)
    ctx = _verified_owner_context(owner)
    owner_sub = stable_owner_sub(ctx)
    canonical_ids = [_canonical_source_id(source_id, owner) for source_id in requested_ids]
    canonical_base = _canonical_source_id(base_file_id, owner) if base_file_id else None
    if canonical_base is not None and canonical_base not in canonical_ids:
        raise ProvenanceRequired("模板来源未包含在生成证明中")
    _validate_excel_write_shape(
        sheet=sheet,
        cells=cells,
        output_name=output_name,
    )
    source_metas = [_load_meta(source_id) for source_id in canonical_ids]
    return artifact_provenance.mint_artifact_evidence(
        owner_sub=owner_sub,
        content_fingerprint_value=artifact_provenance.excel_edit_fingerprint(
            base_file_id=canonical_base,
            sheet=sheet,
            cells=cells,
            output_name=output_name,
        ),
        source_metas=source_metas,
    )


def _owner_of_unchecked(file_id: str) -> str | None:
    """文件创建者（发布时记录的 operated_by）；文件不存在抛 FileError。归属校验用。"""
    fid = _check_id(file_id)
    if _is_legacy_id(fid):
        return None
    return _artifact_meta(fid, require_ready=False).get("operated_by")


def get_download_info(file_id: str, owner: VerifiedArtifactOwner) -> ArtifactDownload:
    """Resolve and integrity-check a ready Artifact without exposing its storage key."""
    fid = _authorized_owner_id(file_id, owner)
    if _is_legacy_id(fid):
        raise ArtifactUnavailable("文件不存在或无权访问", "not_found_or_forbidden")
    meta = _artifact_meta(fid, require_ready=True)
    ext = _ext_of(meta.get("filename", ""))
    if (
        meta.get("filename") != _safe_filename(meta.get("filename", ""))
        or meta.get("storage_key") != _storage_key(fid, ext)
    ):
        raise ArtifactUnavailable("文件元数据校验失败", "metadata_invalid")
    try:
        stored = get_artifact_store().read_bytes(
            meta["storage_key"], max_bytes=_MAX_DOWNLOAD_BYTES
        )
    except ArtifactStoreUnavailable:
        raise
    except ArtifactObjectInvalid as exc:
        reason = (
            "size_limit"
            if exc.reason_code == "object_oversize"
            else exc.reason_code
        )
        message = (
            "文件超过允许的下载大小"
            if reason == "size_limit"
            else "文件对象不存在、不可用或已清理"
        )
        raise ArtifactUnavailable(message, reason) from exc
    except FileError as exc:
        raise ArtifactStoreUnavailable("文件对象状态暂时无法确认") from exc
    if stored.size_bytes != meta.get("size_bytes") or stored.sha256 != meta.get("sha256"):
        raise ArtifactUnavailable("文件完整性校验失败", "integrity_failed")
    return ArtifactDownload(
        artifact_id=fid,
        content=stored.content,
        filename=meta["filename"],
        media_type=_safe_download_media_type(meta),
        size_bytes=meta["size_bytes"],
        sha256=meta["sha256"],
    )
